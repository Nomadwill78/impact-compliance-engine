"""
Multi-Format Document Parser
==============================
Supports: PDF (.pdf) | Word (.docx) | Excel (.xlsx) | Plain text (.txt)
"""
from __future__ import annotations

import io
import re
import time
from pathlib import Path

from app.core.config import settings
from app.models.schemas import FileType, NarrativeSection, ParsedDocument, TableData

_HEADING_PATTERNS = [
    re.compile(r"^([A-Z][A-Z\s&/\-]{4,60})$", re.MULTILINE),
    re.compile(r"^(\d+(?:\.\d+)?\s+[A-Z][^\n]{3,60})$", re.MULTILINE),
    re.compile(r"^((?:[A-Z][a-z]+\s){2,6}[A-Z][a-z]*)$", re.MULTILINE),
]
_KNOWN_SECTIONS = {
    "executive summary", "introduction", "environmental", "social", "governance",
    "financial statements", "financial highlights", "risk factors",
    "materiality assessment", "stakeholder engagement", "climate risk",
    "scope 1", "scope 2", "scope 3", "ghg emissions", "carbon",
    "diversity", "human rights", "supply chain", "audit", "assurance",
    "notes to financial statements", "revenue", "net income",
    "forward-looking statements", "water", "waste", "biodiversity",
    "employee health", "safety", "community investment",
}


def _detect_sections(text: str) -> list[NarrativeSection]:
    lines = text.split("\n")
    heading_indices: list[tuple[int, str]] = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped or len(stripped) > 120:
            continue
        is_heading = any(p.match(stripped) for p in _HEADING_PATTERNS)
        if stripped.lower() in _KNOWN_SECTIONS:
            is_heading = True
        if is_heading:
            heading_indices.append((i, stripped))
    sections: list[NarrativeSection] = []
    for idx, (line_idx, heading) in enumerate(heading_indices):
        next_line = heading_indices[idx + 1][0] if idx + 1 < len(heading_indices) else len(lines)
        body_lines = [ln.strip() for ln in lines[line_idx + 1 : next_line] if ln.strip()]
        body = " ".join(body_lines)
        sections.append(NarrativeSection(
            heading=heading, body=body,
            page_number=max(1, line_idx // 40 + 1),
            word_count=len(body.split()),
        ))
    return sections


def _cap_rows(rows: list, max_rows: int) -> list:
    """Truncate a list of table rows to at most max_rows to bound memory use
    when a document contains an extremely large table or sheet."""
    if len(rows) > max_rows:
        return rows[:max_rows]
    return rows


class _PDFParser:
    @staticmethod
    def parse(file_bytes: bytes, filename: str) -> ParsedDocument:
        import PyPDF2  # noqa: PLC0415

        warnings_list: list[str] = []
        page_texts: list[str] = []
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        for i, page in enumerate(reader.pages):
            try:
                page_texts.append(page.extract_text() or "")
            except Exception as exc:  # noqa: BLE001
                warnings_list.append(f"Page {i + 1}: {exc}")
                page_texts.append("")
        raw_text = "\n".join(page_texts)
        if not raw_text.strip():
            warnings_list.append("No text layer found — document may be image-based.")

        tables: list[TableData] = []
        try:
            import pdfplumber  # noqa: PLC0415

            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for pn, page in enumerate(pdf.pages, 1):
                    for raw in page.extract_tables() or []:
                        if not raw:
                            continue
                        headers = [str(h or "").strip() for h in raw[0]]
                        rows = _cap_rows(
                            [[str(c or "").strip() for c in row] for row in raw[1:]],
                            settings.MAX_TABLE_ROWS,
                        )
                        tables.append(TableData(page_number=pn, headers=headers, rows=rows, row_count=len(rows)))
        except Exception as exc:  # noqa: BLE001
            warnings_list.append(f"Table extraction skipped: {exc}")

        return ParsedDocument(
            filename=filename, file_type=FileType.PDF,
            page_count=len(reader.pages), raw_text=raw_text,
            tables=tables, sections=_detect_sections(raw_text), warnings=warnings_list,
        )


class _DOCXParser:
    @staticmethod
    def parse(file_bytes: bytes, filename: str) -> ParsedDocument:
        try:
            from docx import Document  # noqa: PLC0415
        except ImportError:
            return ParsedDocument(filename=filename, file_type=FileType.DOCX,
                                   warnings=["python-docx not installed."])
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        raw_text = "\n".join(paragraphs)
        tables: list[TableData] = []
        for tbl in doc.tables:
            rows_data = _cap_rows(
                [[cell.text.strip() for cell in row.cells] for row in tbl.rows],
                settings.MAX_TABLE_ROWS,
            )
            if rows_data:
                tables.append(TableData(headers=rows_data[0], rows=rows_data[1:], row_count=len(rows_data) - 1))
        return ParsedDocument(
            filename=filename, file_type=FileType.DOCX,
            page_count=max(1, len(raw_text.split()) // 400),
            raw_text=raw_text, tables=tables, sections=_detect_sections(raw_text),
        )


class _XLSXParser:
    @staticmethod
    def parse(file_bytes: bytes, filename: str) -> ParsedDocument:
        try:
            import openpyxl  # noqa: PLC0415
        except ImportError:
            return ParsedDocument(filename=filename, file_type=FileType.XLSX,
                                   warnings=["openpyxl not installed."])

        # read_only=True streams rows instead of loading the full workbook into
        # memory, which matters a lot for large spreadsheets.
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
        text_parts: list[str] = []
        tables: list[TableData] = []
        warnings_list: list[str] = []
        max_rows = settings.MAX_TABLE_ROWS

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            all_rows = []
            truncated = False
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    all_rows.append(cells)
                if len(all_rows) >= max_rows:
                    truncated = True
                    break
            if truncated:
                warnings_list.append(
                    f"Sheet '{sheet_name}' truncated to {max_rows} rows (MAX_TABLE_ROWS)."
                )
            if not all_rows:
                continue
            tables.append(TableData(sheet_name=sheet_name, headers=all_rows[0], rows=all_rows[1:], row_count=len(all_rows) - 1))
            text_parts.append(f"[Sheet: {sheet_name}]")
            for row in all_rows:
                text_parts.append(" | ".join(c for c in row if c.strip()))

        raw_text = "\n".join(text_parts)
        wb.close()
        return ParsedDocument(
            filename=filename, file_type=FileType.XLSX,
            page_count=len(wb.sheetnames), raw_text=raw_text,
            tables=tables, sections=_detect_sections(raw_text), warnings=warnings_list,
        )


class _TXTParser:
    @staticmethod
    def parse(file_bytes: bytes, filename: str) -> ParsedDocument:
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                raw_text = file_bytes.decode(enc)
                break
            except UnicodeDecodeError:
                continue
        else:
            raw_text = file_bytes.decode("utf-8", errors="replace")
        return ParsedDocument(
            filename=filename, file_type=FileType.TXT,
            page_count=max(1, len(raw_text.split()) // 400),
            raw_text=raw_text, sections=_detect_sections(raw_text),
        )


_EXT_MAP: dict[str, type] = {".pdf": _PDFParser, ".docx": _DOCXParser, ".xlsx": _XLSXParser, ".txt": _TXTParser}
_MIME_MAP: dict[str, type] = {
    "application/pdf": _PDFParser, "application/x-pdf": _PDFParser,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _DOCXParser,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": _XLSXParser,
    "text/plain": _TXTParser,
}


class DocumentParser:
    def parse(self, file_bytes: bytes, filename: str, content_type: str = "") -> ParsedDocument:
        t0 = time.perf_counter()
        ext = Path(filename).suffix.lower()
        parser_cls = _EXT_MAP.get(ext) or _MIME_MAP.get(content_type)
        if parser_cls is None:
            return ParsedDocument(filename=filename, file_type=FileType.UNKNOWN,
                                   warnings=[f"Unsupported file type: '{ext}'. Supported: PDF, DOCX, XLSX, TXT."])
        doc = parser_cls.parse(file_bytes, filename)
        doc.parse_duration_ms = round((time.perf_counter() - t0) * 1000, 2)
        return doc
