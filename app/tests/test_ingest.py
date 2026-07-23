"""
Backend Unit Tests — runs on in-memory SQLite, Gemini API fully mocked.
Run with: python -m pytest app\tests -v --asyncio-mode=auto
"""
from __future__ import annotations

import asyncio
import io
from typing import AsyncGenerator
from unittest.mock import AsyncMock, patch

import pytest
from fastapi.testclient import TestClient
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

from app.core.database import Base, get_db_session
from app.main import app

# ── Test DB ──────────────────────────────────────────────────────────────────
_engine = create_async_engine("sqlite+aiosqlite:///:memory:", connect_args={"check_same_thread": False})
_TestSession = async_sessionmaker(bind=_engine, class_=AsyncSession, expire_on_commit=False, autoflush=False)


async def _override_db() -> AsyncGenerator[AsyncSession, None]:
    async with _TestSession() as session:
        try:
            yield session
            await session.commit()
        except Exception:
            await session.rollback()
            raise


app.dependency_overrides[get_db_session] = _override_db


# ── Fixtures ───────────────────────────────────────────────────────────────��─
@pytest.fixture(scope="session", autouse=True)
def event_loop():
    loop = asyncio.new_event_loop()
    yield loop
    loop.close()


@pytest.fixture(scope="session", autouse=True)
async def setup_db():
    async with _engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    yield
    async with _engine.begin() as conn:
        await conn.run_sync(Base.metadata.drop_all)


@pytest.fixture(scope="module")
def client() -> TestClient:
    return TestClient(app)


@pytest.fixture
def mock_llm():
    """Mocks the combined analyze_document() call used by the /upload endpoint.

    Function-scoped (not module-scoped) so each test gets a fresh mock with
    a clean call count. A module-scoped mock would accumulate call counts
    across every test in the module, making assertions like
    `mock_llm.assert_called_once()` unreliable depending on test order.
    """
    with patch("app.api.endpoints._llm.analyze_document", new_callable=AsyncMock) as m:
        m.return_value = {"entities": [], "compliance": {"gap_summary": "Mocked summary.", "gaps": []}}
        yield m


# ── Helpers ──────────────────────────────────────────────────────────────────
def _pdf(text: str = "Hello World") -> bytes:
    body = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET"
    stream_bytes = body.encode("latin-1")
    stream_len = len(stream_bytes)
    return (
        b"%PDF-1.4\n"
        b"1 0 obj\n<</Type/Catalog/Pages 2 0 R>>\nendobj\n"
        b"2 0 obj\n<</Type/Pages/Count 1/Kids[3 0 R]>>\nendobj\n"
        b"3 0 obj\n<</Type/Page/Parent 2 0 R/Resources<</Font<</F1 4 0 R>>>>/MediaBox[0 0 612 792]/Contents 5 0 R>>\nendobj\n"
        b"4 0 obj\n<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>\nendobj\n"
        b"5 0 obj\n<</Length " + str(stream_len).encode() + b">>\nstream\n" + stream_bytes + b"\nendstream\nendobj\n"
        b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n0000000052 00000 n \n0000000101 00000 n \n0000000208 00000 n \n0000000275 00000 n \n"
        b"trailer\n<</Size 6/Root 1 0 R>>\nstartxref\n356\n%%EOF"
    )


def _txt(text: str) -> bytes:
    return text.encode("utf-8")


def _docx() -> bytes:
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("Scope 1 emissions totalled 1200 tCO2e.")
        doc.add_paragraph("Board of directors: 9 members, 6 independent.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        return b""


def _esg_pdf() -> bytes:
    return _pdf(
        "ESG Report 2024. Scope 1 direct emissions: 1200 tCO2e. Scope 2: 800 tCO2e. "
        "Scope 3 value chain emissions assessed. Water withdrawal: 450 megalitres. "
        "Employee turnover rate 8.5 percent. New hires: 340. "
        "Board of directors: 9 members, 6 independent. Audit committee. "
        "Net-zero target by 2040 validated by SBTi. "
        "CEO statement: sustainability strategy. Materiality assessment. "
        "Independent auditor EY. Forward-looking statements disclaimer."
    )


def _anomalous_pdf() -> bytes:
    return _pdf(
        "Report 2024. Total revenue 100 million USD in Q1. Revenue: 250 million USD. "
        "Market segments: APAC 45%, EMEA 35%, Americas 40%."
    )


# ── Tests ────────────────────────────────────────────────────────────────────
class TestHealth:
    def test_root(self, client: TestClient) -> None:
        r = client.get("/health")
        assert r.status_code == 200
        assert r.json()["status"] == "ok"

    def test_api(self, client: TestClient) -> None:
        r = client.get("/api/v1/health")
        assert r.status_code == 200
        assert r.json()["rules_loaded"] > 0
        assert ".pdf" in r.json()["supported_formats"]


class TestUpload:
    def test_pdf(self, client: TestClient, mock_llm) -> None:
        r = client.post("/api/v1/upload", files={"file": ("test.pdf", _pdf("Compliance test."), "application/pdf")})
        assert r.status_code == 200
        d = r.json()
        assert "report_id" in d
        assert d["file_type"] == "pdf"
        assert 0 <= d["score"] <= 100

    def test_txt(self, client: TestClient, mock_llm) -> None:
        r = client.post("/api/v1/upload", files={"file": ("r.txt", _txt("Board of directors. Scope 1 500 tCO2e."), "text/plain")})
        assert r.status_code == 200
        assert r.json()["file_type"] == "txt"

    def test_docx(self, client: TestClient, mock_llm) -> None:
        d = _docx()
        if not d:
            pytest.skip("python-docx not installed")
        r = client.post("/api/v1/upload",
                         files={"file": ("r.docx", d, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
        assert r.status_code == 200
        assert r.json()["file_type"] == "docx"

    def test_unsupported(self, client: TestClient) -> None:
        r = client.post("/api/v1/upload", files={"file": ("d.csv", b"a,b\n1,2", "text/csv")})
        assert r.status_code == 422
        assert r.json()["detail"]["error"] == "unsupported_file_type"

    def test_empty(self, client: TestClient) -> None:
        r = client.post("/api/v1/upload", files={"file": ("e.pdf", b"", "application/pdf")})
        assert r.status_code == 422
        assert r.json()["detail"]["error"] == "empty_file"

    def test_uses_combined_analyze_document(self, client: TestClient, mock_llm) -> None:
        """Verifies /upload calls the single combined analyze_document() path,
        not the legacy two-call extract_entities()/assess_compliance_gaps() path."""
        r = client.post("/api/v1/upload", files={"file": ("c.pdf", _pdf("Combined call check."), "application/pdf")})
        assert r.status_code == 200
        mock_llm.assert_called_once()

    def test_frameworks_and_sector_params_accepted(self, client: TestClient, mock_llm) -> None:
        r = client.post(
            "/api/v1/upload?frameworks=GRI&frameworks=SASB&sector=education",
            files={"file": ("fs.pdf", _pdf("Framework and sector filter test."), "application/pdf")},
        )
        assert r.status_code == 200
        assert "report_id" in r.json()


class TestCompliance:
    def test_anomalous_has_findings(self, client: TestClient, mock_llm) -> None:
        r = client.post("/api/v1/upload", files={"file": ("a.pdf", _anomalous_pdf(), "application/pdf")})
        assert r.status_code == 200
        assert r.json()["finding_count"] > 0

    def test_esg_scores_higher(self, client: TestClient, mock_llm) -> None:
        bare = client.post("/api/v1/upload", files={"file": ("b.pdf", _pdf("Nothing."), "application/pdf")}).json()["score"]
        rich = client.post("/api/v1/upload", files={"file": ("e.pdf", _esg_pdf(), "application/pdf")}).json()["score"]
        assert rich >= bare

    def test_gri_disclosed(self, client: TestClient, mock_llm) -> None:
        r = client.post("/api/v1/upload", files={"file": ("g.pdf", _esg_pdf(), "application/pdf")})
        assert r.json()["gri_disclosed_count"] > 0


class TestRetrieval:
    def test_get_report(self, client: TestClient, mock_llm) -> None:
        up = client.post("/api/v1/upload", files={"file": ("r.pdf", _pdf("Retrieve."), "application/pdf")})
        rid = up.json()["report_id"]
        r = client.get(f"/api/v1/reports/{rid}")
        assert r.status_code == 200
        body = r.json()
        assert body["report_id"] == rid
        assert "findings" in body
        assert "gri_disclosures" in body
        assert "sasb_metrics" in body

    def test_404(self, client: TestClient) -> None:
        r = client.get("/api/v1/reports/00000000-0000-0000-0000-000000000000")
        assert r.status_code == 404

    def test_list(self, client: TestClient, mock_llm) -> None:
        client.post("/api/v1/upload", files={"file": ("l.pdf", _pdf("List."), "application/pdf")})
        r = client.get("/api/v1/reports?limit=5")
        assert r.status_code == 200
        assert isinstance(r.json(), list)


class TestAnalytics:
    def test_summary(self, client: TestClient, mock_
